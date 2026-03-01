import streamlit as st
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
import os
from dotenv import load_dotenv
import pypdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time

# ================= 1. SYSTEM CONFIGURATION =================
st.set_page_config(
    page_title="Nova Assistant: Elite Admissions", 
    page_icon="🌠", 
    layout="wide",
    initial_sidebar_state="expanded"
)
load_dotenv()

# --- CSS Styling ---
st.markdown("""
<style>
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        font-weight: bold;
        background-color: #1E3A8A;
        color: white;
    }
    .stButton>button:hover {
        background-color: #2563EB;
        color: white;
    }
    h1 { color: #1E3A8A; }
    h2 { color: #1E3A8A; font-size: 1.5rem; }
    .reportview-container { background: #f0f2f6; }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d1fae5;
        color: #065f46;
        border: 1px solid #34d399;
    }
    .supp-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #eff6ff;
        border: 1px solid #bfdbfe;
        color: #1e3a8a;
    }
</style>
""", unsafe_allow_html=True)

# --- API Key Setup ---
try:
    api_key = st.secrets.get("GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        st.warning("⚠️ API Key not found. Please set GOOGLE_API_KEY in .env or Streamlit Secrets.")
    else:
        genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"Setup Error: {e}")

# ================= 2. ELITE LOGIC SPECIFICATIONS =================

# I. PERSONAL STATEMENT LOGIC
FULL_ESSAY_SPEC = """
LOGIC SPEC: Common App Personal Statement (Strict 550–600 Words)

A) NON-NEGOTIABLES
- Target word count: 575 words (Range: 550-600).
- Core requirement: Reveal how the student thinks/chooses, not just what they did.
- Max 2 full scenes total.
- NO Moralizing: Do not use "I learned," "This taught me," "I realized." Show change through behavior.
- NO Motivational Ending: No "change the world," "dream come true."

B) INPUTS & LENS SELECTION
- Define a "Lens": A repeatable pattern of operation.
- Format: "When [trigger] happens, I tend to [instinct], so I [action]."
- Valid only if proven in 2 different contexts.

C) STRUCTURE: THEMATIC NARRATIVE (Default)
1. Hook (55-80 words): Lens in action. Immediate motion. No quotes.
2. Scene 1 (170-210 words): Origin moment. Must include TENSION and an A/B DECISION POINT.
   - Format: "I could do A... or B... So I chose B."
3. Scene 2 (170-210 words): Later moment where lens becomes intentional. Must include Internal Processing (cognition, not emotion).
4. Closing (55-80 words): Grounded. How they operate now. "Operating Manual" style.

D) VOICE & STYLE
- Tone: Sparky, simple, precise.
- Uniqueness Device: Choose ONE (System Thinking, Translation, Signal Detection, or Controlled Stubbornness).
- NO generic transitions (Moreover, Furthermore).
- NO Trauma Dumping.

E) SCENE WRITING SPEC
- Setup -> Tension -> Decision Point (A/B) -> Internal Processing -> Action -> Consequence.
- Connection line: Links scene to lens without moralizing.

F) IMPROVISATION CONTEXT (Uzbekistan)
- If details are missing, improvise using local context:
  - Zakovat (Intellectual games)
  - Math Olympiads (National obsession)
  - Mahalla (Community duties)
  - Lyceum vs. Public School dynamics
  - Shadow Education (Training centers/Repetitors)
"""

# II. RECOMMENDATION LOGIC
FULL_REC_SPEC = """
LOGIC SPEC: Recommendation Letter (Counselor + Teachers + Director)

A) NON-NEGOTIABLES
- Authenticity over poetry. Sounds like an adult educator.
- EVIDENCE BASED: Every claim backed by a specific story/measurable detail.
- NO Resume repeating. Select 2-4 main proof points.
- Include 1 "Growth Edge" (small weakness that becomes strength).

B) ROLES
1. SCHOOL DIRECTOR (Counselor Role): Focus on community impact, leadership, maturity, school context.
2. TEACHER (Subject Specific): Focus on intellectual habits, classroom behavior, "how they learn."

C) UZBEKISTAN CONTEXT
- Directors often act as counselors.
- Teachers emphasize discipline, respect, and national curriculum (Prezident maktabi, specialized subjects).
- Mentions of class rank are common and acceptable.

D) STRUCTURE
1. Opening: Credibility + Relationship + One-line summary.
2. Academic Strength: A vivid classroom moment/story.
3. Character/Community: How they elevate peers.
4. Growth Edge: Safe weakness + response.
5. Closing: Strong endorsement.

E) IMPROVISATION ENGINE
- If details missing, generate plausible:
  - Classroom debates (e.g., Alisher Navoi literature analysis, Physics lab on optics).
  - Leadership moments (Class monitor, cleaning day organizer - Hashar).
  - Peer tutoring.
- DO NOT invent specific awards/test scores unless inferred.
"""

# III. ACTIVITIES LOGIC
FULL_ACTIVITIES_SPEC = """
LOGIC SPEC: Common App Activities List
1. Analyze record. If activities < 10, IMPROVISE up to 6 high-quality entries fitting an ambitious Uzbek student.
2. Potential Improvised Activities:
   - English Tutor (Private lessons)
   - Family Business Helper (Bazaar/Shop logistics)
   - Mahalla Volunteer (Community aid)
   - Sports (Football, Kurash, Chess)
   - Telegram Channel Admin (Educational blog)

3. STRICT OUTPUT FORMAT (Repeat for each activity):
   Activity type: [Select from Common App types]
   Position/Leadership: [Max 50 chars]
   Organization Name: [Max 100 chars]
   Description: [Max 150 chars - use action verbs, impact, no full sentences]
   Participation grade levels: [e.g., 9, 10, 11]
   Timing of participation: [During school year / During school break / All year]
   Hours spent per week: [Number]
   Weeks spent per year: [Number]
"""

# IV. SUPPLEMENTAL ESSAY LOGIC (DETAILED)
SUPPLEMENTAL_LOGIC_SPEC = """
LOGIC SPEC: SUPPLEMENTAL ESSAYS (Analyze & Execute)

TASK:
1. Analyze the user's prompt to classify it into one of the 8 types below.
2. Use the 'Perfect answer logic' for that specific type.
3. USE YOUR INTERNAL KNOWLEDGE BASE to find real labs, professors, values, or traditions for the specific University.
4. Integrate the Student Record + Generated Personal Statement to ensure consistency.
5. APPLY THE "ELITE WRITING STYLE" defined below strictly.

ELITE WRITING STYLE (The "Nova" Voice):
[... Keep your existing logic here ...]
"""

# ================= 3. HELPER FUNCTIONS =================

def extract_text_from_pdf(file):
    """Robust PDF text extraction."""
    try:
        pdf_reader = pypdf.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text if text.strip() else None
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def call_gemini(system_prompt, user_content, temperature=0.7):
    """Wrapper for Gemini generation."""
    full_prompt = f"{system_prompt}\n\nINPUT DATA:\n{user_content}"
    
    try:
        # FIXED: Changed from retired preview model to stable version
        model = genai.GenerativeModel("gemini-2.5-flash")
        
        response = model.generate_content(
            contents=[{"role": "user", "parts": [full_prompt]}],
            generation_config=GenerationConfig(temperature=temperature, max_output_tokens=4000)
        )
        return response.text
    except Exception as e:
        return f"Error generating content: {e}"

def create_docx_report(data_dict, supplemental_list=None):
    """Generates a professional DOCX file with all content."""
    doc = Document()
    
    # Style Setup
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    head = doc.add_heading(f"Nova Admissions Report: {data_dict.get('name', 'Student')}", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d')}")
    doc.add_paragraph("Confidential Application Materials").italic = True
    doc.add_page_break()

    # SECTION 1: ANALYSIS
    doc.add_heading('1. Student Analysis & Lens Strategy', level=1)
    doc.add_paragraph(data_dict.get('analysis', 'N/A'))
    doc.add_page_break()

    # SECTION 2: ACTIVITIES
    doc.add_heading('2. Extracurricular Activities (Common App)', level=1)
    doc.add_paragraph(data_dict.get('activities', 'N/A'))
    doc.add_page_break()

    # SECTION 3: PERSONAL STATEMENT
    doc.add_heading('3. Personal Statement (550-600 Words)', level=1)
    doc.add_paragraph("Formatting note: Ensure paragraphs are separated by blank lines.").italic = True
    doc.add_paragraph(data_dict.get('essay', 'N/A'))
    doc.add_page_break()

    # SECTION 4: DIRECTOR REC
    doc.add_heading('4. School Director / Counselor Recommendation', level=1)
    doc.add_paragraph(data_dict.get('director_rec', 'N/A'))
    doc.add_page_break()

    # SECTION 5: TEACHER 1 REC
    doc.add_heading('5. Teacher Recommendation 1', level=1)
    doc.add_paragraph(data_dict.get('teacher1_rec', 'N/A'))
    doc.add_page_break()

    # SECTION 6: TEACHER 2 REC
    doc.add_heading('6. Teacher Recommendation 2', level=1)
    doc.add_paragraph(data_dict.get('teacher2_rec', 'N/A'))

    # SECTION 7: SUPPLEMENTALS
    if supplemental_list:
        doc.add_page_break()
        doc.add_heading('7. Supplemental Essays', level=1)
        for idx, supp in enumerate(supplemental_list, 1):
            doc.add_heading(f"Supplement {idx}: {supp['university']}", level=2)
            doc.add_paragraph(f"Prompt: {supp['prompt']}")
            doc.add_paragraph(f"Word Limit: {supp.get('word_limit', 'N/A')}")
            doc.add_paragraph("-" * 20)
            doc.add_paragraph(supp['content'])
            doc.add_paragraph("-" * 40)
            doc.add_page_break()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= 4. MAIN APPLICATION FLOW =================

def main():
    st.title("🌠 Nova Assistant")
    st.markdown("### Elite Admissions & Supplemental Strategist")
    
    # --- SESSION STATE INITIALIZATION ---
    if "step" not in st.session_state: st.session_state.step = 1
    if "student_text" not in st.session_state: st.session_state.student_text = ""
    if "analysis_result" not in st.session_state: st.session_state.analysis_result = ""
    if "student_name" not in st.session_state: st.session_state.student_name = "Student"
    if "generated_content" not in st.session_state: st.session_state.generated_content = {}
    if "supplementals" not in st.session_state: st.session_state.supplementals = []

    # --- SIDEBAR ---
    with st.sidebar:
        st.header("1. Upload Record")
        uploaded_file = st.file_uploader("Upload Student PDF", type=["pdf"])
        
        if uploaded_file and st.session_state.step == 1:
            if st.button("Start Processing"):
                with st.spinner("Extracting text and analyzing record..."):
                    text = extract_text_from_pdf(uploaded_file)
                    if text:
                        st.session_state.student_text = text
                        
                        # Quick Analysis
                        analysis_prompt = """
                        Analyze this student record. 
                        1. Identify the full Student Name.
                        2. Identify intended major.
                        3. Identify the names of Teacher 1 and Teacher 2 if listed.
                        4. Identify School Director Name if listed.
                        Output: Plain text summary.
                        """
                        analysis = call_gemini(analysis_prompt, text)
                        st.session_state.analysis_result = analysis
                        
                        try:
                            lines = analysis.split('\n')
                            for line in lines:
                                if "Name" in line and ":" in line:
                                    st.session_state.student_name = line.split(":")[1].strip()
                                    break
                        except:
                            pass
                            
                        st.session_state.step = 2
                        st.rerun()
                    else:
                        st.error("Could not extract text. File might be an image-only PDF.")

    # --- MAIN PAGE LOGIC ---

    # STEP 1: WELCOME
    if st.session_state.step == 1:
        st.info("👋 Upload a student PDF record (Uzbekistan Format) to begin.")
        st.markdown("""
        **System Capabilities:**
        * **Core Application**: Personal Statement, Activities, Recommendations.
        * **NEW: Supplemental Strategist**: Analyzes specific college prompts (Why Us, Why Major) and writes targeted essays.
        """)

    # STEP 2: REVIEW & GENERATE
    elif st.session_state.step == 2:
        st.success("✅ Record Analyzed")
        col1, col2 = st.columns([1, 1])
        with col1:
            st.subheader("Profile Summary")
            st.write(st.session_state.analysis_result)
        with col2:
            st.subheader("Configuration")
            st.info("Ready to generate Core Application.")
            improv_mode = st.radio("Improvisation Mode", ["Balanced (Recommended)", "Creative (High Improvisation)"])
            
            if st.button("🚀 Generate Core Application"):
                st.session_state.step = 3
                st.session_state.improv_mode = improv_mode
                st.rerun()
            if st.button("Start Over"):
                st.session_state.clear()
                st.rerun()

    # STEP 3 & 4: GENERATION PROCESS & SUPPLEMENTALS
    elif st.session_state.step >= 3:
        if not st.session_state.generated_content:
            st.subheader(f"Generating Core Application for {st.session_state.student_name}...")
            
            # 1. Activities
            with st.spinner("1/4 Generating Activities List..."):
                prompt = FULL_ACTIVITIES_SPEC
                if st.session_state.improv_mode == "Creative":
                    prompt += "\nNOTE: Improvise 6 distinct activities to fill gaps."
                res = call_gemini(prompt, st.session_state.student_text)
                st.session_state.generated_content["activities"] = res

            # 2. Personal Statement
            with st.spinner("2/4 Writing Personal Statement..."):
                res = call_gemini(FULL_ESSAY_SPEC, st.session_state.student_text, temperature=0.8)
                st.session_state.generated_content["essay"] = res

            # 3. Recommendations
            with st.spinner("3/4 Writing Recommendations..."):
                prompt_d = FULL_REC_SPEC + "\nROLE: SCHOOL DIRECTOR (Counselor Context). Focus on leadership and community."
                res_d = call_gemini(prompt_d, st.session_state.student_text)
                st.session_state.generated_content["director_rec"] = res_d
                
                prompt_t1 = FULL_REC_SPEC + "\nROLE: TEACHER 1 (Subject Specific). Focus on academic rigor."
                res_t1 = call_gemini(prompt_t1, st.session_state.student_text)
                st.session_state.generated_content["teacher1_rec"] = res_t1

                prompt_t2 = FULL_REC_SPEC + "\nROLE: TEACHER 2 (Different Subject). Focus on class dynamics."
                res_t2 = call_gemini(prompt_t2, st.session_state.student_text)
                st.session_state.generated_content["teacher2_rec"] = res_t2
            
            st.rerun()

        st.success("🎉 Core Application Generated")
        core_tabs = st.tabs(["Essay", "Activities", "Director Rec", "Teacher 1", "Teacher 2"])
        
        with core_tabs[0]: st.markdown(st.session_state.generated_content["essay"])
        with core_tabs[1]: st.text(st.session_state.generated_content["activities"])
        with core_tabs[2]: st.markdown(st.session_state.generated_content["director_rec"])
        with core_tabs[3]: st.markdown(st.session_state.generated_content["teacher1_rec"])
        with core_tabs[4]: st.markdown(st.session_state.generated_content["teacher2_rec"])

        st.markdown("---")

        # ================= STEP 4: SUPPLEMENTAL ESSAY GENERATOR =================
        st.markdown("## 📚 Supplemental Essay Generator")
        st.markdown("""
        <div class="supp-box">
        <b>Advanced Feature:</b> This tool analyzes the specific "secret" logic of the prompt and builds bridges to your profile.
        </div>
        """, unsafe_allow_html=True)
        
        c1, c2, c3 = st.columns([2, 3, 1])
        with c1:
            univ_name = st.text_input("University Name", placeholder="e.g. Duke University")
        with c2:
            supp_prompt = st.text_area("Essay Prompt", placeholder="e.g. Why do you want to study here?")
        with c3:
            word_limit = st.text_input("Word Limit", placeholder="e.g. 250 words")

        if st.button("✨ Generate Supplemental Essay"):
            if not univ_name or not supp_prompt:
                st.error("Please provide University Name and Prompt.")
            else:
                with st.spinner(f"🔍 Analyzing prompt logic for {univ_name}..."):
                    context_data = f"""
                    STUDENT RECORD: {st.session_state.student_text}
                    ALREADY GENERATED PERSONAL STATEMENT: {st.session_state.generated_content["essay"]}
                    ALREADY GENERATED ACTIVITIES: {st.session_state.generated_content["activities"]}
                    TARGET UNIVERSITY: {univ_name}
                    ESSAY PROMPT: {supp_prompt}
                    WORD LIMIT: {word_limit}
                    """
                    
                    supp_result = call_gemini(
                        system_prompt=SUPPLEMENTAL_LOGIC_SPEC, 
                        user_content=context_data
                    )
                    
                    st.session_state.supplementals.append({
                        "university": univ_name,
                        "prompt": supp_prompt,
                        "word_limit": word_limit,
                        "content": supp_result
                    })
                    st.rerun()

        if st.session_state.supplementals:
            st.subheader("Generated Supplementals")
            for i, supp in enumerate(st.session_state.supplementals):
                with st.expander(f"{supp['university']} - {supp['prompt'][:50]}...", expanded=True):
                    st.markdown(supp['content'])
                    if st.button(f"Delete Essay {i+1}", key=f"del_{i}"):
                        st.session_state.supplementals.pop(i)
                        st.rerun()

        # FINAL DOWNLOAD
        final_data = {
            "name": st.session_state.student_name,
            "analysis": st.session_state.analysis_result,
            "activities": st.session_state.generated_content["activities"],
            "essay": st.session_state.generated_content["essay"],
            "director_rec": st.session_state.generated_content["director_rec"],
            "teacher1_rec": st.session_state.generated_content["teacher1_rec"],
            "teacher2_rec": st.session_state.generated_content["teacher2_rec"]
        }
        
        docx_file = create_docx_report(final_data, st.session_state.supplementals)
        st.download_button(
            label="📄 Download Complete Application Package",
            data=docx_file,
            file_name=f"Nova_App_Package_{st.session_state.student_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
        if st.button("Start New Student"):
            st.session_state.clear()
            st.rerun()

if __name__ == "__main__":
    main()
